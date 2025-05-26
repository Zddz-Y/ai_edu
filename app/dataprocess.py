import re
import uuid
import subprocess
import pathlib
from docx import Document
import os
from docx.oxml.ns import nsmap, qn
from docx.text.paragraph import Paragraph

def iter_textbox_paragraphs(doc):
    """
    生成器：遍历整个文档，找出 w:txbxContent 里的 w:p，
    yield 回 python-docx 的 Paragraph 对象。
    """
    # 1) 找到所有 <w:txbxContent> 节点
    for txbx in doc._element.xpath('.//w:txbxContent', namespaces=nsmap):
        # 2) 其中的每个 <w:p> 当作独立段落返回
        for p in txbx.xpath('.//w:p', namespaces=nsmap):
            yield Paragraph(p, doc)


def save_run_image(image_part, qid):
    """保存图片并返回文件名"""
    fmt = image_part.content_type.split("/")[-1]  # png / jpeg / x-wmf
    uid = uuid.uuid4().hex[:8]
    fname = f"{qid}_{uid}.{fmt}"
    out_dir = pathlib.Path(r"E:\NLP_Model\ai_edu\data\processed_data\images")
    out_dir.mkdir(exist_ok=True)
    path = out_dir / fname
    path.write_bytes(image_part.blob)
    # 如果是 WMF，再转 PNG 备份
    if fmt == "x-wmf":
        png = out_dir / f"{qid}_{uid}.png"
        try:
            subprocess.run(
                ["magick", "convert", str(path), str(png)],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                check=True,
            )
            print(f"已将 WMF 转换为 PNG: {png}")
            # 删除原始 WMF 文件
            os.remove(path)
            print(f"已删除原始 WMF 文件: {path}")
            return png.name  # 返回 PNG 文件名
        except subprocess.CalledProcessError as e:
            print(f"WMF 转换失败: {e}")
            return fname  # 返回原始 WMF 文件名
    print("-------------------已运行save_run_image方法---------------------")
    return fname

def extract_images_from_runs(paragraph, doc, qid):
    images = []
    segment = ""
    for run in paragraph.runs:
        xml = run._element.xml
        text = run.text
        # 文字
        if text:
            segment += text
        # 图片
        
        rids = re.findall(r'(?:r:id|r:embed)="(rId\d+)"', xml)
        for rid in rids:
            try:
                # rel = doc.part.rels.get(rid)
                # print(f"处理 rId: {rid}, target_ref: {rel.target_ref}")  # 调试信息
                if rid in doc.part.related_parts:
                    # 处理图片
                    part = doc.part.related_parts[rid]
                    if part.content_type.startswith("image/"):
                        fname = save_run_image(part, qid)
                        images.append(fname)
                        segment += f"[IMG:{fname}]"
            except Exception as e:
                print(f"处理rId {rid}时出错: {e}")
    return images, segment

def extract_doc_content(doc_path):
    """提取文档中的文本、图片及结构化题目，并在content中按原位置插入图片占位符"""
    doc = Document(doc_path)
    questions = []
    current = None
    collecting = False

    # 获取文档中所有元素的顺序（段落和表格）
    def get_document_elements():
        """获取文档中段落和表格的顺序"""
        elements = []
        for element in doc.element.body:
            if element.tag.endswith('}p'):  # 段落
                # 找到对应的段落对象
                for para in doc.paragraphs:
                    if para._element == element:
                        elements.append(('paragraph', para))
                        break
            elif element.tag.endswith('}tbl'):  # 表格
                # 找到对应的表格对象
                for table in doc.tables:
                    if table._element == element:
                        elements.append(('table', table))
                        break
        return elements


    for element_type, element in get_document_elements():
        if element_type == 'paragraph':
            para = element
            # 先检查本段是否新题号
            raw = para.text.strip()
            m_q = re.match(r'^(\d+)\.', raw)
            if m_q:
                # 关闭上一个题
                if current:
                    questions.append(current)
                qid = m_q.group(1)
                # 新题初始化
                current = {
                    "id": int(qid),
                    "content": "",     # 将按 run 拼接
                    "options": [],
                    "images": [],
                    "formulas": []
                }
                collecting = True

            if not collecting or current is None:
                continue  # 跳过题目前的文字

            # 按 run 遍历，构造本段 content 片段
            images, segment = extract_images_from_runs(para, doc, current["id"])
            current["images"].extend(images)

            # 去除纯空白段
            if not segment.strip():
                continue

            # 判断是否为选项行
            m_opt = re.match(r'^[A-D]\.', segment.strip())
            if m_opt:
                # 新增一个选项 dict，也可以保留占位符在 option text 中
                current["options"].append({
                    "text": segment.strip(),
                    "images": re.findall(r'\[IMG:(.*?)\]', segment)
                })
            else:
                # 普通题干，直接累加，并保留换行
                if current["content"]:
                    current["content"] += "\n" + segment
                else:
                    current["content"] = segment
        elif element_type == 'table':
            table = element
            if collecting and current is not None:
                # 提取表格内容并格式化
                table_content = "\n[表格开始]\n"
                for row_idx, row in enumerate(table.rows):
                    row_cells = []
                    for cell in row.cells:
                        cell_text = ""
                        # 处理单元格中的段落和图片
                        for cell_para in cell.paragraphs:
                            images, cell_segment = extract_images_from_runs(cell_para, doc, current["id"])
                            if images:
                                current["images"].extend(images)
                            if cell_segment.strip():
                                cell_text += cell_segment.strip() + " "
                        row_cells.append(cell_text.strip())
                    table_content += " | ".join(row_cells) + "\n"
                table_content += "[表格结束]\n"
                
                # 将表格内容添加到当前题目的content中
                if current["content"]:
                    current["content"] += table_content
                else:
                    current["content"] = table_content


    # 最后一题
    if current:
        questions.append(current)

    return questions

if __name__ == "__main__":
    qlist = extract_doc_content(r"E:\NLP_Model\ai_edu\data\math\2023年江苏省泰州市中考数学真题（原卷版）.docx")
    

    # 4. 保存结构化数据
    import json
    output_dir = r"E:\NLP_Model\ai_edu\data\processed_data"
    with open(os.path.join(output_dir, "taizhou2023_questions.json"), "w", encoding="utf-8") as f:
        json.dump(qlist, f, indent=2, ensure_ascii=False)